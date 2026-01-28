#!/usr/bin/env python3
"""
Ứng dụng soạn giáo án STEM )
Chạy cục bộ: python3 app.py
Sau đó mở http://localhost:5000
"""

from flask import Flask, render_template, request, jsonify, send_file
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os

app = Flask(__name__)

# Cấu hình API key
API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyBAXeNa1aKD5Re0TIj1ktF_4iVDLAXRbic")
genai.configure(api_key=API_KEY)
model = genai.GenerativeModel('gemini-2.5-flash')

def create_docx(text, title):
    """Tạo file Word từ text"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    
    # Tiêu đề
    hdr = doc.add_heading(f"KẾ HOẠCH BÀI DẠY STEM: {title.upper()}", 0)
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for line in text.split('\n'):
        if line.strip().startswith('###'):
            p = doc.add_heading(line.replace('###', '').strip(), level=1)
        elif line.strip().startswith('##'):
            p = doc.add_heading(line.replace('##', '').strip(), level=2)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/generate', methods=['POST'])
def generate():
    """API endpoint để sinh giáo án"""
    try:
        data = request.json
        ten_bai = data.get('ten_bai', '')
        khoi = data.get('khoi', '6')
        chu_trinh = data.get('chu_trinh', '')
        hoat_dong = data.get('hoat_dong', [])
        
        if not ten_bai:
            return jsonify({'error': 'Vui lòng nhập tên bài dạy'}), 400
        
        prompt_content = f"""
        Bạn là chuyên gia giáo dục STEM tại Việt Nam. Hãy soạn kế hoạch bài dạy bài '{ten_bai}' lớp {khoi}.
        Yêu cầu tuân thủ cấu trúc Công văn 3089/BGDĐT-GDTrH:
        1. Mục tiêu (Kiến thức; Năng lực; Phẩm chất).
        2. Thiết bị dạy học và học liệu.
        3. Tiến trình dạy học: Sử dụng chu trình {chu_trinh}.
        
        Chi tiết các hoạt động: {', '.join(hoat_dong) if hoat_dong else 'Tất cả'}. 
        Mỗi hoạt động PHẢI có đủ 4 mục nhỏ:
        - a) Mục tiêu
        - b) Nội dung
        - c) Sản phẩm
        - d) Tổ chức thực hiện (Giao nhiệm vụ -> Thực hiện -> Báo cáo -> Kết luận).
        Yêu cầu cấu trúc bài dạy:



Mục tiêu bài học: Nêu rõ về kiến thức (Toán học là trọng tâm), kĩ năng, thái độ và năng lực đặc thù (năng lực giải quyết vấn đề, năng lực mô hình hóa toán học).

Thiết bị dạy học và học liệu: Liệt kê cụ thể.

Tiến trình dạy học (5 hoạt động theo CV 3089):

HĐ 1: Xác định vấn đề: Giao nhiệm vụ thực tiễn dẫn đến nhu cầu giải quyết bằng toán học.

HĐ 2: Nghiên cứu kiến thức nền và đề xuất giải pháp: Học sinh tìm hiểu kiến thức toán học liên quan để giải quyết vấn đề.

HĐ 3: Lựa chọn giải pháp/Thiết kế sản phẩm: Học sinh thảo luận, vẽ bản vẽ kỹ thuật hoặc lập kế hoạch tính toán.

HĐ 4: Chế tạo mẫu, thử nghiệm và thảo luận: Thực hiện tính toán/chế tạo và điều chỉnh.

HĐ 5: Chia sẻ, thảo luận và đánh giá: Thuyết trình về sản phẩm và ứng dụng toán học trong đó.
        Văn phong: Sư phạm chuẩn mực, trình bày rõ ràng bằng Markdown.
        """
        
        response = model.generate_content(prompt_content)
        result_text = response.text
        
        return jsonify({
            'success': True,
            'content': result_text,
            'ten_bai': ten_bai
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/download', methods=['POST'])
def download():
    """Download file Word"""
    try:
        data = request.json
        content = data.get('content', '')
        ten_bai = data.get('ten_bai', 'giao_an')
        
        docx_file = create_docx(content, ten_bai)
        
        return send_file(
            docx_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"Giao_an_STEM_{ten_bai}.docx"
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    print("\n" + "="*60)
    print("🎓 Ứng dụng soạn giáo án STEM 3089 )")
    print("="*60)
    print("\n✨ Ứng dụng đang chạy tại: http://localhost:5001")
    print("\n💡 Mẹo:")
    print("   - Nếu port 5001 đã dùng, đổi port: app.run(port=5002)")
    print("   - Thiết lập API key: export GEMINI_API_KEY='your_key_here'")
    print("\n" + "="*60 + "\n")
    
    app.run(debug=True, port=5001, host='0.0.0.0')
